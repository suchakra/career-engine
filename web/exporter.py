"""Multi-format résumé export for the tailored résumé (Markdown / PDF / DOCX).

Turns a parsed :class:`web.tailor.TailoredResume` into downloadable bytes in the
format the user wants:
- **Markdown** — ATS-friendly plain text, copy-paste anywhere.
- **PDF** — polished, via WeasyPrint + an autoescaped HTML template (same engine
  as the master résumé renderer; no browser needed).
- **DOCX** — editable Word doc (recruiters often want to tweak), via python-docx.

The content is format-agnostic (a summary + a list of achievements), so each
renderer is a pure function of the ``TailoredResume`` — no model calls, no
Streamlit. HTML is autoescaped so model-derived text can't inject markup.
"""

from __future__ import annotations

import io

from jinja2 import Environment, select_autoescape

from web.tailor import TailoredResume, tailored_to_markdown

__all__ = [
    "tailored_to_docx_bytes",
    "tailored_to_markdown",
    "tailored_to_pdf_bytes",
]

# Autoescaped HTML template (default=True escapes EVERY variable — model output is
# rendered as inert text, never markup). Print-optimized, single-column, ATS-safe.
# The environment + compiled template are module-level (constant), so the UI path
# doesn't rebuild/recompile them on every call.
_PDF_TEMPLATE_SRC = """<!DOCTYPE html>
<html lang="en"><head><meta charset="UTF-8" />
<title>{{ name or "Tailored résumé" }}</title>
<style>
  * { box-sizing: border-box; margin: 0; padding: 0; }
  body { font-family: "Georgia", "Times New Roman", serif; font-size: 11pt;
         line-height: 1.45; color: #1a1a1a; padding: 2.2cm 2cm; }
  h1 { font-size: 20pt; margin-bottom: 2pt; }
  h2 { font-size: 13pt; border-bottom: 1px solid #999; margin: 16pt 0 8pt;
       padding-bottom: 2pt; text-transform: uppercase; letter-spacing: 0.5pt; }
  .summary { margin-top: 6pt; }
  .ach { margin-bottom: 10pt; }
  .headline { font-weight: bold; }
  .note { font-style: italic; color: #555; font-size: 10pt; }
</style></head><body>
  <h1>{{ name or "Tailored résumé" }}</h1>
  {% if summary %}<p class="summary">{{ summary }}</p>{% endif %}
  {% if achievements %}<h2>Selected achievements</h2>
  {% for a in achievements %}<div class="ach">
    <div class="headline">{{ a.headline }}</div>
    {% if a.full_text %}<div>{{ a.full_text }}</div>{% endif %}
    {% if a.relevance_note %}<div class="note">Why it fits: {{ a.relevance_note }}</div>{% endif %}
  </div>{% endfor %}{% endif %}
</body></html>
"""

_ENV = Environment(autoescape=select_autoescape(default=True, default_for_string=True))
_PDF_TEMPLATE = _ENV.from_string(_PDF_TEMPLATE_SRC)  # compiled once at import


def tailored_to_pdf_bytes(tailored: TailoredResume, *, name: str = "") -> bytes:
    """Render a tailored résumé to PDF bytes (WeasyPrint; autoescaped HTML)."""
    html_str = _PDF_TEMPLATE.render(
        name=name, summary=tailored.summary, achievements=tailored.achievements
    )
    import weasyprint

    pdf_bytes: bytes = weasyprint.HTML(string=html_str).write_pdf()
    return pdf_bytes


def tailored_to_docx_bytes(tailored: TailoredResume, *, name: str = "") -> bytes:
    """Render a tailored résumé to editable DOCX bytes (python-docx)."""
    from docx import Document

    doc = Document()
    doc.add_heading(name or "Tailored résumé", level=0)
    if tailored.summary:
        doc.add_paragraph(tailored.summary)
    if tailored.achievements:
        doc.add_heading("Selected achievements", level=1)
        for a in tailored.achievements:
            bullet = doc.add_paragraph(style="List Bullet")
            bullet.add_run(a.headline).bold = True
            if a.full_text:
                doc.add_paragraph(a.full_text)
            if a.relevance_note:
                note = doc.add_paragraph()
                note.add_run(f"Why it fits: {a.relevance_note}").italic = True

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
