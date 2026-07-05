"""Render a :class:`web.resume_builder.StructuredResume` to Markdown / PDF / DOCX.

Real, ATS-safe layout: contact header · summary · skills · experience grouped by
role (title — org · dates, quantified bullets) · education. No internal
"why it fits" reasoning appears in the document. Each renderer is a pure function
of the résumé; HTML is autoescaped so model-derived text can't inject markup.
"""

from __future__ import annotations

import io

from jinja2 import Environment, select_autoescape

from web.resume_builder import StructuredResume

__all__ = [
    "structured_to_docx_bytes",
    "structured_to_markdown",
    "structured_to_pdf_bytes",
]


def _contact_line(resume: StructuredResume) -> str:
    c = resume.contact
    parts = [c.email, c.phone, c.location, *c.links]
    return " · ".join(p for p in parts if p.strip())


def structured_to_markdown(resume: StructuredResume) -> str:
    """Render the résumé as ATS-friendly Markdown."""
    c = resume.contact
    lines: list[str] = [f"# {c.name or 'Résumé'}"]
    contact = _contact_line(resume)
    if contact:
        lines += ["", contact]
    if resume.summary:
        lines += ["", "## Summary", "", resume.summary]
    if resume.skills:
        lines += ["", "## Skills", "", " · ".join(resume.skills)]
    if resume.experience:
        lines += ["", "## Experience"]
        for role in resume.experience:
            header = " — ".join(p for p in (role.title, role.org) if p)
            lines += ["", f"### {header}"]
            if role.dates:
                lines.append(f"*{role.dates}*")
            lines += [f"- {b}" for b in role.bullets]
    if resume.education:
        lines += ["", "## Education"]
        for role in resume.education:
            header = " — ".join(p for p in (role.title, role.org) if p)
            date_suffix = f" ({role.dates})" if role.dates else ""
            lines.append(f"- {header}{date_suffix}")
    return "\n".join(lines).rstrip() + "\n"


_PDF_TEMPLATE_SRC = """<!DOCTYPE html>
<html lang="en"><head><meta charset="UTF-8" /><title>{{ c.name or "Résumé" }}</title>
<style>
  * { box-sizing: border-box; margin: 0; padding: 0; }
  body { font-family: "Georgia","Times New Roman",serif; font-size: 10.5pt;
         line-height: 1.4; color: #1a1a1a; padding: 1.8cm 1.9cm; }
  header { text-align: center; margin-bottom: 10pt; }
  h1 { font-size: 20pt; }
  .contact { font-size: 9.5pt; color: #444; margin-top: 3pt; }
  h2 { font-size: 12pt; text-transform: uppercase; letter-spacing: 0.6pt;
       border-bottom: 1px solid #888; margin: 14pt 0 6pt; padding-bottom: 2pt; }
  .role { margin-top: 8pt; }
  .role-head { font-weight: bold; }
  .dates { float: right; font-weight: normal; color: #555; font-size: 9.5pt; }
  ul { margin: 3pt 0 0 16pt; }
  li { margin-bottom: 2pt; }
  .skills { }
</style></head><body>
  <header>
    <h1>{{ c.name or "Résumé" }}</h1>
    {% if contact %}<div class="contact">{{ contact }}</div>{% endif %}
  </header>
  {% if r.summary %}<h2>Summary</h2><p>{{ r.summary }}</p>{% endif %}
  {% if r.skills %}<h2>Skills</h2><p class="skills">{{ r.skills | join(" · ") }}</p>{% endif %}
  {% if r.experience %}<h2>Experience</h2>
  {% for role in r.experience %}<div class="role">
    <div class="role-head">{{ [role.title, role.org] | select | join(" — ") }}
      {% if role.dates %}<span class="dates">{{ role.dates }}</span>{% endif %}</div>
    {% if role.bullets %}<ul>{% for b in role.bullets %}<li>{{ b }}</li>{% endfor %}</ul>{% endif %}
  </div>{% endfor %}{% endif %}
  {% if r.education %}<h2>Education</h2>
  {% for role in r.education %}<div class="role">
    <div class="role-head">{{ [role.title, role.org] | select | join(" — ") }}
      {% if role.dates %}<span class="dates">{{ role.dates }}</span>{% endif %}</div>
  </div>{% endfor %}{% endif %}
</body></html>
"""

_ENV = Environment(autoescape=select_autoescape(default=True, default_for_string=True))
_PDF_TEMPLATE = _ENV.from_string(_PDF_TEMPLATE_SRC)


def structured_to_pdf_bytes(resume: StructuredResume) -> bytes:
    """Render the résumé to PDF bytes (WeasyPrint; autoescaped HTML)."""
    html_str = _PDF_TEMPLATE.render(r=resume, c=resume.contact, contact=_contact_line(resume))
    import weasyprint

    pdf_bytes: bytes = weasyprint.HTML(string=html_str).write_pdf()
    return pdf_bytes


def structured_to_docx_bytes(resume: StructuredResume) -> bytes:
    """Render the résumé to editable DOCX bytes (python-docx)."""
    from docx import Document

    doc = Document()
    doc.add_heading(resume.contact.name or "Résumé", level=0)
    contact = _contact_line(resume)
    if contact:
        doc.add_paragraph(contact)
    if resume.summary:
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(resume.summary)
    if resume.skills:
        doc.add_heading("Skills", level=1)
        doc.add_paragraph(" · ".join(resume.skills))
    if resume.experience:
        doc.add_heading("Experience", level=1)
        for role in resume.experience:
            head = " — ".join(p for p in (role.title, role.org) if p)
            para = doc.add_paragraph()
            para.add_run(head).bold = True
            if role.dates:
                para.add_run(f"   {role.dates}").italic = True
            for bullet in role.bullets:
                doc.add_paragraph(bullet, style="List Bullet")
    if resume.education:
        doc.add_heading("Education", level=1)
        for role in resume.education:
            head = " — ".join(p for p in (role.title, role.org) if p)
            suffix = f" ({role.dates})" if role.dates else ""
            doc.add_paragraph(f"{head}{suffix}", style="List Bullet")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
